import random
import time
from docx import Document
from variables import document_path, input_text

def append_text_to_doc(doc_path, paragraphs):
    try:
        # Load the existing document
        doc = Document(doc_path)
    except Exception as e:
        print(f"Error loading document: {e}")
        return

    for paragraph_text in paragraphs:
        if not paragraph_text.strip():
            continue

        words = paragraph_text.split()
        total_words = len(words)
        appended_words = 0

        while appended_words < total_words:
            # Randomly select the number of words to append (between 10 and 15)
            word_count = random.randint(15, 30)
            chunk = words[appended_words:appended_words + word_count]

            if not chunk:
                break

            # Append the chunk as a new paragraph if it's the first chunk, otherwise continue the last paragraph
            if appended_words == 0:
                doc.add_paragraph(' '.join(chunk))
            else:
                doc.paragraphs[-1].add_run(' ' + ' '.join(chunk))

            appended_words += len(chunk)
            print(f"Appended {len(chunk)} words. Total appended: {appended_words}/{total_words}")

            # Save the document after each update
            try:
                doc.save(doc_path)
            except Exception as e:
                print(f"Error saving document: {e}")
                return

            # Wait for 60 seconds before appending the next chunk
            wait_time = 15
            print(f"Waiting for {wait_time} seconds before the next append...")
            time.sleep(wait_time)

    print("Text appending completed.")

# Example usage
# document_path = "/mnt/c/Users/mjfie/OneDrive/Documents/FakeVersionHistory/Test3.docx"  # Path to the Word document
# input_text = [
#     "Beyond its effects on performance, creatine indirectly promotes muscle hypertrophy by enabling users to lift heavier weights and complete more repetitions. Additionally, creatine draws water into muscle cells, increasing cell volume. This process, called cell volumization, may signal anabolic pathways and contribute to muscle growth over time. Moreover, creatine’s ability to reduce muscle damage and inflammation following intense workouts can significantly speed up recovery. It also helps replenish phosphocreatine stores more quickly post-exercise, preparing the body for subsequent sessions. While its primary benefits are seen in physical performance, creatine also supports cognitive function. Research suggests that it may improve mental clarity and reduce fatigue, especially during tasks requiring high focus or problem-solving.",
#     "One of the most compelling reasons for people who work out to use creatine is its proven safety and efficacy. Creatine monohydrate is one of the safest and most scientifically validated supplements available. Decades of research have confirmed its effectiveness in improving strength, muscle mass, and performance, with minimal risks when used as directed. Its benefits are not limited to elite athletes; creatine can enhance the performance of recreational gym-goers and those new to fitness. By enabling users to train harder and recover faster, it serves as a versatile supplement that aligns with a wide range of fitness goals."
# ]

append_text_to_doc(document_path, input_text)
